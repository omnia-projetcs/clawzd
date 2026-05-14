"""
Clawzd — Research Export Pipeline.
Handles PDF, DOCX, and PPTX export with:
- Mermaid diagram → PNG conversion via Playwright
- Proper Table of Contents generation
- Full UTF-8 Unicode support
- Quality dashboard header
- Professional styling
"""
import os
import re
import json
import asyncio
import logging
import hashlib
from datetime import datetime, timezone

logger = logging.getLogger("clawzd.research.export")


# ── Mermaid → PNG Conversion ─────────────────────────────────────────────────

async def mermaid_to_png(mermaid_code: str, output_path: str, width: int = 800) -> bool:
    """Convert a Mermaid diagram to a PNG image using Playwright.
    
    Returns True on success, False on failure.
    """
    try:
        from playwright.async_api import async_playwright
    except ImportError:
        logger.warning("Playwright not installed — skipping mermaid-to-png")
        return False

    # Build a minimal HTML page that renders the mermaid diagram
    html_content = f"""<!DOCTYPE html>
<html><head>
<script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
<style>body {{ background: white; margin: 20px; font-family: sans-serif; }}</style>
</head><body>
<pre class="mermaid">{mermaid_code}</pre>
<script>mermaid.initialize({{ startOnLoad: true, theme: 'default' }});</script>
</body></html>"""

    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page(viewport={"width": width, "height": 600})
            await page.set_content(html_content)
            # Wait for mermaid to render
            await page.wait_for_timeout(2000)
            # Find the rendered SVG
            svg_el = await page.query_selector("svg.mermaid")
            if not svg_el:
                svg_el = await page.query_selector("pre.mermaid svg")
            if not svg_el:
                # Fallback: screenshot the whole body
                svg_el = await page.query_selector("body")
            if svg_el:
                await svg_el.screenshot(path=output_path, type="png")
                await browser.close()
                logger.info("Mermaid → PNG: %s", output_path)
                return True
            await browser.close()
    except Exception as e:
        logger.warning("mermaid_to_png failed: %s", e)
    return False


def _extract_mermaid_blocks(md: str) -> list[tuple[str, str]]:
    """Extract mermaid code blocks from markdown.
    
    Returns list of (full_match, mermaid_code) tuples.
    """
    pattern = re.compile(r'```mermaid\s*\n([\s\S]*?)```', re.MULTILINE)
    return [(m.group(0), m.group(1).strip()) for m in pattern.finditer(md)]


def sanitize_mermaid_code(code: str) -> str:
    """Clean up mermaid code for rendering — strip HTML tags, fix labels."""
    # Strip HTML <br> tags
    code = re.sub(r'<br\s*/?>', ' ', code, flags=re.IGNORECASE)
    # Auto-quote labels with special chars in square brackets
    code = re.sub(
        r'(\w+)\[([^\]"]+[()/<>][^\]"]*)\]',
        lambda m: f'{m.group(1)}["{m.group(2).strip()}"]',
        code,
    )
    # Auto-quote labels with special chars in parentheses (for round nodes)
    code = re.sub(
        r'(\w+)\(([^)"]+[\[\]/<>][^)"]*)\)',
        lambda m: f'{m.group(1)}("{m.group(2).strip()}")',
        code,
    )
    # Auto-quote labels with special chars in curly braces (for diamond/decision nodes)
    code = re.sub(
        r'(\w+)\{([^}"]+[()[\]/<>][^}"]*)\}',
        lambda m: f'{m.group(1)}{{"{m.group(2).strip()}"}}',
        code,
    )
    return code


async def replace_mermaid_with_images(md: str, assets_dir: str) -> str:
    """Replace all mermaid blocks in markdown with PNG image references.
    
    Used for PDF/DOCX exports where mermaid JS rendering isn't available.
    """
    blocks = _extract_mermaid_blocks(md)
    if not blocks:
        return md

    for full_match, code in blocks:
        clean_code = sanitize_mermaid_code(code)
        code_hash = hashlib.md5(clean_code.encode()).hexdigest()[:8]
        img_name = f"mermaid_{code_hash}.png"
        img_path = os.path.join(assets_dir, img_name)

        if not os.path.exists(img_path):
            success = await mermaid_to_png(clean_code, img_path)
            if not success:
                # Leave the code block as-is but clean it
                md = md.replace(full_match, f"```\n{clean_code}\n```")
                continue

        # Replace with image reference
        md = md.replace(full_match, f"![Diagram]({img_path})")

    return md


# ── Markdown → HTML Conversion ───────────────────────────────────────────────

def _md_to_html(md: str, title: str = "Research Report") -> str:
    """Convert markdown to styled HTML for PDF rendering."""
    try:
        import markdown as _md
        html_body = _md.markdown(
            md,
            extensions=["tables", "fenced_code", "toc", "nl2br"],
        )
    except ImportError:
        # Manual fallback
        html_body = _manual_md_to_html(md)

    return f"""<!DOCTYPE html>
<html lang="fr"><head>
<meta charset="UTF-8">
<title>{_esc(title)}</title>
<style>{_get_pdf_css()}</style>
</head><body>
{html_body}
</body></html>"""


def _manual_md_to_html(md: str) -> str:
    """Basic manual markdown → HTML for when the markdown lib isn't available."""
    import html as _html
    lines = md.split("\n")
    out = []
    in_code = False
    in_table = False
    for line in lines:
        if line.startswith("```"):
            if in_code:
                out.append("</code></pre>")
                in_code = False
            else:
                lang = line[3:].strip()
                out.append(f'<pre><code class="language-{lang}">')
                in_code = True
            continue
        if in_code:
            out.append(_html.escape(line))
            continue
        # Tables
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                out.append("<table>")
                in_table = True
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                continue  # separator row
            cells = [c.strip() for c in line.split("|") if c.strip()]
            tag = "th" if not any("<td>" in o for o in out[-5:]) else "td"
            out.append("<tr>" + "".join(f"<{tag}>{c}</{tag}>" for c in cells) + "</tr>")
            continue
        elif in_table:
            out.append("</table>")
            in_table = False
        # Headings
        if line.startswith("### "):
            out.append(f"<h3>{line[4:]}</h3>")
        elif line.startswith("## "):
            out.append(f"<h2>{line[3:]}</h2>")
        elif line.startswith("# "):
            out.append(f"<h1>{line[2:]}</h1>")
        elif line.startswith("- "):
            out.append(f"<li>{line[2:]}</li>")
        elif line.startswith("> "):
            out.append(f"<blockquote>{line[2:]}</blockquote>")
        elif line.strip() == "---":
            out.append("<hr>")
        elif line.strip():
            out.append(f"<p>{line}</p>")
        else:
            out.append("<br>")
    if in_table:
        out.append("</table>")
    return "\n".join(out)


def _esc(text: str) -> str:
    import html as _html
    return _html.escape(text)


def _get_pdf_css() -> str:
    """Professional CSS for PDF export."""
    return """
    @page { size: A4; margin: 2cm; }
    body {
        font-family: 'Segoe UI', 'Helvetica Neue', Arial, sans-serif;
        font-size: 11pt; line-height: 1.6; color: #1a1a2e;
        max-width: 100%;
    }
    h1 { font-size: 22pt; color: #1a1a2e; border-bottom: 3px solid #6366f1;
         padding-bottom: 8px; margin-top: 24px; }
    h2 { font-size: 16pt; color: #312e81; margin-top: 20px;
         border-bottom: 1px solid #e5e7eb; padding-bottom: 4px; }
    h3 { font-size: 13pt; color: #4338ca; margin-top: 16px; }
    table { border-collapse: collapse; width: 100%; margin: 12px 0; }
    th, td { border: 1px solid #d1d5db; padding: 8px 12px; text-align: left; font-size: 10pt; }
    th { background: #f3f4f6; font-weight: 600; color: #374151; }
    tr:nth-child(even) { background: #f9fafb; }
    blockquote { border-left: 4px solid #6366f1; padding: 8px 16px;
                 margin: 12px 0; background: #f8f9ff; color: #4b5563; font-style: italic; }
    code { background: #f3f4f6; padding: 2px 6px; border-radius: 4px;
           font-family: 'JetBrains Mono', monospace; font-size: 9pt; }
    pre { background: #1e1e2e; color: #cdd6f4; padding: 16px; border-radius: 8px;
          overflow-x: auto; font-size: 9pt; }
    pre code { background: none; padding: 0; color: inherit; }
    img { max-width: 100%; border-radius: 4px; margin: 8px 0; }
    a { color: #4f46e5; text-decoration: underline; }
    hr { border: none; border-top: 1px solid #e5e7eb; margin: 24px 0; }
    ul, ol { padding-left: 24px; }
    li { margin-bottom: 4px; }
    .dashboard { background: linear-gradient(135deg, #eef2ff, #ecfdf5);
                 border: 1px solid #d1d5db; border-radius: 8px;
                 padding: 16px; margin-bottom: 20px; display: flex;
                 gap: 24px; align-items: center; }
    .dash-score { font-size: 36pt; font-weight: 800; }
    .dash-score.high { color: #059669; }
    .dash-score.mid { color: #d97706; }
    .dash-score.low { color: #dc2626; }
    .dash-metrics { display: flex; gap: 16px; flex-wrap: wrap; }
    .dash-metric { text-align: center; min-width: 80px; }
    .dash-metric-value { font-size: 18pt; font-weight: 700; color: #1e1b4b; }
    .dash-metric-label { font-size: 8pt; color: #6b7280; text-transform: uppercase; }
    """


def _build_dashboard_html(proj: dict) -> str:
    """Build an HTML quality dashboard for PDF/DOCX export."""
    score = proj.get("current_score", 0)
    pct = round(score * 100)
    iters = len(proj.get("iterations", []))
    sources = len(proj.get("search_results", []))
    assets = len(proj.get("assets", []))
    sc_class = "high" if pct >= 70 else ("mid" if pct >= 40 else "low")

    return f"""<div class="dashboard">
    <div class="dash-score {sc_class}">{pct}%</div>
    <div class="dash-metrics">
        <div class="dash-metric"><div class="dash-metric-value">{sources}</div><div class="dash-metric-label">Sources</div></div>
        <div class="dash-metric"><div class="dash-metric-value">{iters}</div><div class="dash-metric-label">Iterations</div></div>
        <div class="dash-metric"><div class="dash-metric-value">{assets}</div><div class="dash-metric-label">Assets</div></div>
    </div>
    </div>"""


# ── PDF Export ───────────────────────────────────────────────────────────────

async def export_pdf(proj: dict, report_md: str, pdir: str) -> str:
    """Export a research report to PDF. Returns the output file path."""
    assets_dir = os.path.join(pdir, "assets")
    os.makedirs(assets_dir, exist_ok=True)

    # 1. Sanitize and convert mermaid to images
    processed_md = await replace_mermaid_with_images(report_md, assets_dir)

    # 2. Strip __TABLE__/__CHART__ markers — render them as plain text
    processed_md = _strip_structured_markers(processed_md)

    # 3. Build HTML
    title = proj.get("title", "Research Report")
    dashboard_html = _build_dashboard_html(proj)
    html_body = _md_to_html(processed_md, title)
    # Inject dashboard after <body>
    html_body = html_body.replace("<body>\n", f"<body>\n{dashboard_html}\n", 1)

    path = os.path.join(pdir, "report.pdf")

    # Try WeasyPrint first, fallback to FPDF
    try:
        from weasyprint import HTML
        await asyncio.to_thread(lambda: HTML(string=html_body).write_pdf(path))
        logger.info("PDF exported via WeasyPrint: %s", path)
        return path
    except ImportError:
        logger.info("WeasyPrint not available — using fpdf2 fallback")
    except Exception as e:
        logger.warning("WeasyPrint failed: %s — trying fpdf2 fallback", e)

    # FPDF2 fallback with UTF-8 font
    try:
        from fpdf import FPDF
        class UTF8PDF(FPDF):
            pass
        
        def _gen_pdf_fallback():
            pdf = UTF8PDF()
            pdf.add_page()
            # Try to use a Unicode font
            try:
                pdf.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", uni=True)
                pdf.add_font("DejaVu", "B", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", uni=True)
                pdf.set_font("DejaVu", "B", 16)
            except Exception:
                pdf.set_font("Helvetica", "B", 16)

            pdf.cell(0, 10, title[:80], ln=True)
            # Score header
            score = proj.get("current_score", 0)
            try:
                pdf.set_font("DejaVu", "", 10)
            except Exception:
                pdf.set_font("Helvetica", "", 10)
            pdf.cell(0, 8, f"Quality: {score:.0%} | Sources: {len(proj.get('search_results', []))} | Iterations: {len(proj.get('iterations', []))}", ln=True)
            pdf.ln(4)

            for line in processed_md.split("\n"):
                if line.startswith("# "):
                    try:
                        pdf.set_font("DejaVu", "B", 14)
                    except Exception:
                        pdf.set_font("Helvetica", "B", 14)
                    pdf.cell(0, 8, line[2:][:100], ln=True)
                elif line.startswith("## "):
                    try:
                        pdf.set_font("DejaVu", "B", 12)
                    except Exception:
                        pdf.set_font("Helvetica", "B", 12)
                    pdf.cell(0, 7, line[3:][:100], ln=True)
                elif line.startswith("### "):
                    try:
                        pdf.set_font("DejaVu", "B", 11)
                    except Exception:
                        pdf.set_font("Helvetica", "B", 11)
                    pdf.cell(0, 7, line[4:][:100], ln=True)
                elif line.strip().startswith("![") and "](" in line:
                    # Image reference — try to embed
                    img_match = re.search(r'!\[.*?\]\((.+?)\)', line)
                    if img_match:
                        img_path = img_match.group(1)
                        if os.path.isfile(img_path):
                            try:
                                pdf.image(img_path, w=160)
                            except Exception:
                                pass
                elif line.strip():
                    try:
                        pdf.set_font("DejaVu", "", 10)
                    except Exception:
                        pdf.set_font("Helvetica", "", 10)
                    # Handle encoding gracefully
                    try:
                        pdf.multi_cell(0, 5, line)
                    except Exception:
                        safe = line.encode("latin-1", "replace").decode("latin-1")
                        pdf.multi_cell(0, 5, safe)
            pdf.output(path)

        await asyncio.to_thread(_gen_pdf_fallback)
        logger.info("PDF exported via fpdf2: %s", path)
        return path
    except ImportError:
        raise ImportError("Neither weasyprint nor fpdf2 is installed")


# ── DOCX Export ──────────────────────────────────────────────────────────────

async def export_docx(proj: dict, report_md: str, pdir: str) -> str:
    """Export a research report to DOCX. Returns the output file path."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    assets_dir = os.path.join(pdir, "assets")
    os.makedirs(assets_dir, exist_ok=True)

    # Convert mermaid to images
    processed_md = await replace_mermaid_with_images(report_md, assets_dir)
    # Strip structured markers
    processed_md = _strip_structured_markers(processed_md)

    path = os.path.join(pdir, "report.docx")

    def _gen_docx():
        doc = Document()
        title = proj.get("title", "Research Report")
        doc.add_heading(title, 0)

        # Dashboard paragraph
        score = proj.get("current_score", 0)
        iters = len(proj.get("iterations", []))
        sources = len(proj.get("search_results", []))
        p = doc.add_paragraph()
        run = p.add_run(f"Quality: {score:.0%}  |  Sources: {sources}  |  Iterations: {iters}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # spacing

        in_code = False
        code_lines = []

        for line in processed_md.split("\n"):
            if line.startswith("```"):
                if in_code:
                    # End code block
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(8)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    code_lines = []
                    in_code = False
                else:
                    in_code = True
                continue

            if in_code:
                code_lines.append(line)
                continue

            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            elif line.strip().startswith("![") and "](" in line:
                # Image — try to embed
                img_match = re.search(r'!\[.*?\]\((.+?)\)', line)
                if img_match:
                    img_path = img_match.group(1)
                    if os.path.isfile(img_path):
                        try:
                            doc.add_picture(img_path, width=Inches(5.5))
                        except Exception:
                            doc.add_paragraph(f"[Image: {img_path}]")
                    else:
                        doc.add_paragraph(f"[Image: {img_path}]")
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.runs[0] if p.runs else p.add_run("")
                run.font.italic = True
                run.font.color.rgb = RGBColor(80, 80, 120)
            elif line.strip() == "---":
                doc.add_paragraph("_" * 60)
            elif line.strip():
                doc.add_paragraph(line)

        doc.save(path)

    await asyncio.to_thread(_gen_docx)
    logger.info("DOCX exported: %s", path)
    return path


# ── Helpers ──────────────────────────────────────────────────────────────────

def _strip_structured_markers(md: str) -> str:
    """Remove __TABLE__/__CHART__ markers and render them as plain text."""
    # Convert __TABLE__ to markdown tables
    def _table_to_md(match):
        try:
            config = json.loads(match.group(1))
            headers = config.get("headers", [])
            rows = config.get("rows", [])
            title = config.get("title", "")
            lines = []
            if title:
                lines.append(f"**{title}**\n")
            if headers:
                lines.append("| " + " | ".join(headers) + " |")
                lines.append("| " + " | ".join("---" for _ in headers) + " |")
            for row in rows:
                lines.append("| " + " | ".join(str(c) for c in row) + " |")
            return "\n".join(lines)
        except Exception:
            return ""

    md = re.sub(r'__TABLE__(\{[\s\S]*?\})__TABLE__', _table_to_md, md)
    md = re.sub(r'__CHART__\{[\s\S]*?\}__CHART__', '', md)
    md = re.sub(r'__PROGRESS__\{[\s\S]*?\}__PROGRESS__', '', md)
    md = re.sub(r'__CARD__\{[\s\S]*?\}__CARD__', '', md)
    md = re.sub(r'__ALERT__\{[\s\S]*?\}__ALERT__', '', md)
    md = re.sub(r'__ARTIFACT__\{[\s\S]*?\}__ARTIFACT__', '', md)
    return md
